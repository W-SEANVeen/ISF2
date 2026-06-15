"""Generate readme.docx — project structure and file/folder descriptions."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# ── title ───────────────────────────────────────────────────────────────
title = doc.add_heading('《孤城》VR 游戏项目 — 文件结构说明', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('项目路径：e:\\GameProject\\7_VR_test\\Test')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('本文档说明项目各主要目录及文件的内容与作用。', style='Normal')

# ════════════════════════════════════════════════════════════════════════
# 1. 顶层目录
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('一、顶层目录与根文件', level=1)

top_level = [
    ('Assets/', 'Unity 项目核心资源目录，存放所有游戏资源（场景、脚本、模型、材质、音效等），是开发中主要工作目录。'),
    ('ProjectSettings/', 'Unity 项目设置目录，包含图形、输入、物理、播放器等全局配置（.asset 文件），由 Unity Editor 自动管理。'),
    ('Packages/', 'Unity Package Manager 包清单目录，记录项目所依赖的官方与第三方包（manifest.json）。'),
    ('Docs/', '项目文档目录，存放实验报告、PPT、交互逻辑文档等开发与答辩资料。'),
    ('UserSettings/', 'Unity Editor 的用户级设置（编辑器布局、搜索设置等），不参与版本构建。'),
    ('Library/', 'Unity 导入缓存目录，保存资源处理结果（自动生成，可删除后由 Editor 重建）。'),
    ('Logs/', 'Unity Editor 运行时日志输出目录。'),
    ('Temp/', 'Unity 临时文件目录，构建与导入过程的临时数据（可安全删除）。'),
    ('PICO-Unity-Integration-SDK-release_3.1.0/', 'PICO VR 头显的原生 Unity 集成 SDK（3.1.0 版），包含平台 API、空间音频、企业功能等模块。'),
    ('TestUnity Live Preview Plugin-1.0.5-20250211/', 'PICO Live Preview 插件，用于在 PC Editor 中实时预览 PICO VR 设备画面，加快开发迭代。'),
]

for name, desc in top_level:
    h = doc.add_heading('', level=2)
    run = h.add_run(name)
    run.font.size = Pt(12)
    doc.add_paragraph(desc, style='Normal')

# ════════════════════════════════════════════════════════════════════════
# 2. Assets 子目录
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('二、Assets 子目录详解', level=1)

assets_dirs = [
    ('Scenes/', 'Unity 场景文件。包含主游戏场景（如 Scene2_city_v2.unity），是游戏世界的容器，放置所有 GameObject。'),
    ('Script/', 'C# 脚本目录，按功能子文件夹分类组织。'),
    ('Prefab/', '预制体（Prefab）目录，存放可复用的游戏对象模板（如箭矢、云朵、UI 面板等）。'),
    ('Animation/', '动画剪辑文件（.anim），包含角色攻击、死亡等动作片段（如死士 Attack.anim、骑兵 Dead.anim）。'),
    ('animator/', 'Animator Controller 文件（.controller），管理角色状态机与动画过渡逻辑。'),
    ('model/', '3D 模型文件（.fbx）及部分 Prefab，包括角色（士兵、马匹）、建筑（塔楼、城门、房屋）、武器（弩、箭矢）、地形（terrain_magou）等。'),
    ('material/', '基础材质文件（.mat），定义颜色、贴图等表面属性。'),
    ('Materials/', '另一些材质文件（如雪花粒子材质 snowflake.mat、snowdot.mat）。'),
    ('shader/', 'Shader 与材质目录，包含 Shader Graph（.shadergraph）和手写着色器（.shader），实现水墨描边（PropOutline）、地形、飘旗等视觉特效。'),
    ('Textures/', '纹理贴图文件（.png/.jpg），包括水墨笔触、宣纸背景、山峦贴图、雪花粒子等，为水墨画风格提供视觉素材。'),
    ('Audio/', '音效资源。SFX 子目录存放战争号角、射箭破空声、鼓声等战场音效（.mp3/.wav）。'),
    ('Fonts/', '字体文件，包括 HarmonyOS Sans SC（鸿蒙字体）和庞门正道粗书体两种中文字体，及 TextMesh Pro SDF 配置。'),
    ('VFX/', 'Visual Effect Graph 粒子效果文件（.vfx），用于实现下雪等场景特效。'),
    ('Resources/', 'Unity Resources 可加载资源目录（可通过 Resources.Load() 运行时加载的资源）。'),
    ('ScriptableObject/', 'ScriptableObject 配置资源，如 AudioCollectionSO（音效集合配置）。'),
    ('Settings/', '项目自定义设置资源。'),
    ('Plugins/', '原生插件目录（Android 平台 AndroidManifest 等）。'),
    ('XR/', 'XR 子系统配置，包含 PICO XR Loader 和运行设置（PXR_Loader.asset、PXR_Settings.asset）。'),
    ('XRI/', 'XR Interaction Toolkit 编辑器设置（XRInteractionEditorSettings）。'),
    ('TextMesh Pro/', 'TextMesh Pro 资源（字体、着色器、精灵、样式表），用于高品质文字渲染。'),
    ('TutorialInfo/', 'URP 模板自带的 Tutorial 脚本与图标（Readme.cs）。'),
    ('pic/', 'UI 与背景图片，如水墨风按钮背景、远山剪影、城墙贴图等。'),
    ('rig/', '角色骨骼绑定（Rigging）相关资源。'),
    ('Animation.meta 等 .meta 文件', 'Unity 资源导入元数据文件，每个资源对应一个 .meta，记录 GUID 与导入设置，由 Unity 自动维护，不可删除。'),
]

for name, desc in assets_dirs:
    h = doc.add_heading('', level=2)
    run = h.add_run(name)
    run.font.size = Pt(12)
    doc.add_paragraph(desc, style='Normal')

# ════════════════════════════════════════════════════════════════════════
# 3. Script 目录详解
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('三、Script 脚本目录详解', level=1)

doc.add_paragraph(
    '所有 C# 脚本位于 Assets/Script/ 下，按模块分子目录，以下为各脚本的角色说明：'
)

script_sections = [
    ('managers /  管理器模块', [
        ('GameManager.cs', '游戏全局管理器，控制游戏状态（进行中、胜利、失败）与场景流程。'),
        ('SceneFlowManager.cs', '场景流程控制器，管理场景切换逻辑与过渡。'),
        ('BattleDirector.cs', '战斗导演控制器，驱动战场事件序列（如敌方波次进攻时间线）。'),
        ('SiegeManager.cs', '攻城战管理器，控制攻城塔、云梯等 siege 事件的触发与推进。'),
        ('EventManager.cs', '事件管理器，提供全局事件发布/订阅机制，解耦模块间通信。'),
        ('AudioManager.cs', '音效管理器，控制背景音乐与 SFX 的播放、切换和音量。'),
        ('PlayerHealth.cs', '玩家生命值管理，扣血、死亡判定与 UI 更新。'),
        ('HornPlayer.cs', '号角音效播放器，触发战争号角等氛围音效。'),
        ('BootLoader.cs', '场景启动加载器，初始化必要的管理器与服务。'),
        ('SceneChanger.cs', '场景切换组件，处理异步场景加载与过渡动画。'),
        ('GameEndUIManager.cs', '游戏结束界面 UI 管理器（胜利/失败面板）。'),
        ('MenuUIManager.cs', '主菜单 UI 管理器。'),
        ('XRUIManager.cs', 'XR UI 管理器，处理 VR 界面交互（如射线点击）。'),
    ]),
    ('armymanage /  军队管理模块', [
        ('SquadCommander.cs', '小队指挥官，控制一组士兵的移动、攻击与阵型。'),
        ('ArmyMover.cs', '军队移动控制器，处理单位寻路与编队前进。'),
        ('EnemyAssault.cs', '敌方突击逻辑，控制敌人冲锋与冲击行为。'),
        ('EnemyHorse.cs', '敌方骑兵控制，处理马匹移动与骑乘战斗。'),
        ('EnemyRandomizer.cs', '敌方随机生成器，为敌人外观与行为增加随机变化。'),
        ('ArmyBeanScatterer.cs', '军队散开/散布逻辑，使单位分布更自然。'),
        ('ArrowFlight.cs', '箭矢飞行逻辑，处理弓箭弹道与命中检测。'),
        ('ArrowRainVolley.cs', '箭雨齐射逻辑，控制多支箭矢同时发射。'),
        ('SiegeLadder.cs', '攻城梯控制，管理云梯的展开与搭靠城墙。'),
        ('PorterRetreat.cs', '搬运工撤退逻辑，非战斗单位在适当时机撤离。'),
        ('MicroMotion.cs', '角色细微动作控制器，添加呼吸、摇摆等细节动画。'),
    ]),
    ('obj /  可交互对象模块', [
        ('CrossbowShoot.cs', '弩射击逻辑，控制玩家弩箭发射与装填。'),
        ('ArmyAdvance.cs', '军队推进触发器，检测玩家操作触发的推进事件。'),
        ('Igniter.cs', '点火器，控制火焰/火箭的点火效果。'),
        ('RickIgniteController.cs', '火箭点火控制器，管理火箭点燃与发射流程。'),
    ]),
    ('VR /  VR 交互模块', [
        ('ResetXROrigin.cs', '重置 XR Origin 位置，处理 VR 空间重新定位。'),
        ('StickyGrabInteractable.cs', '粘性抓取交互组件，使物体被抓取后吸附在手上。'),
        ('VRFaceHUD.cs', 'VR 面向 HUD，使 UI 始终跟随玩家视线方向。'),
    ]),
    ('tools /  工具模块', [
        ('GridGenerator.cs', '网格生成器工具，用于场景中辅助网格的生成。'),
        ('AudioCollectionSO.cs', 'ScriptableObject 音效集合定义，组织分组音效数据。'),
    ]),
    ('根目录脚本', [
        ('EnemyCharge.cs', '敌方冲锋逻辑，控制单体敌人冲向目标。'),
        ('TestArrowShooter.cs', '箭矢发射测试脚本，用于调试弓箭系统。'),
        ('TorchFireController.cs', '火炬火焰控制器，管理火焰燃烧效果。'),
        ('WeaponTrailController.cs', '武器拖尾轨迹控制器，实现挥砍时的刀光效果。'),
        ('control_ball.cs', '球体控制脚本（辅助调试或 UI 组件控制）。'),
    ]),
]

for section_title, scripts in script_sections:
    h = doc.add_heading('', level=2)
    run = h.add_run(section_title)
    run.font.size = Pt(12)
    table = doc.add_table(rows=1, cols=2, style='Light Grid Accent 1')
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.5)
    hdr = table.rows[0].cells
    hdr[0].text = '文件名'
    hdr[1].text = '作用'
    for filename, desc in scripts:
        row = table.add_row().cells
        row[0].text = filename
        row[1].text = desc

# ════════════════════════════════════════════════════════════════════════
# 4. 场景文件
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('四、场景文件', level=1)

doc.add_paragraph(
    '场景文件位于 Assets/Scenes/，是 Unity 编辑器中搭建的游戏世界容器，'
    '包含所有场景中的 GameObject、灯光、摄像机与引用资源。'
    '主要场景 Scene2_city_v2.unity 为城市战场主关卡。'
)

# ════════════════════════════════════════════════════════════════════════
# 5. SDK 与插件目录
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('五、SDK 与第三方插件', level=1)

sdk_items = [
    ('PICO-Unity-Integration-SDK-release_3.1.0/', 'PICO VR 官方 Unity SDK（3.1.0），包含：Platform（平台账号、支付、社交）、Runtime（追踪、渲染、输入）、SpatialAudio（空间音频）、Enterprise（企业级功能）等模块，是 VR 功能的基础依赖。'),
    ('TestUnity Live Preview Plugin-1.0.5-20250211/', 'PICO Live Preview 插件，允许开发者在 Unity Editor 中实时预览 PICO VR 设备的画面和交互，无需反复部署到真机。包含 Editor 和 Runtime 两部分。'),
]

for name, desc in sdk_items:
    h = doc.add_heading('', level=2)
    run = h.add_run(name)
    run.font.size = Pt(12)
    doc.add_paragraph(desc)

# ════════════════════════════════════════════════════════════════════════
# 6. Docs 目录
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('六、Docs 文档目录', level=1)

doc.add_paragraph(
    'Docs/ 目录存放项目开发与答辩相关的文档资料，具体包括：'
)

docs_files = [
    ('实验报告.md', '项目实验报告，记录开发过程、技术实现与测试结果。'),
    ('交互逻辑项目进程文档.md', '交互逻辑设计与开发进度追踪文档。'),
    ('项目汇报文档.md', '项目汇报/进度汇报文本。'),
    ('《孤城》VR游戏设计答辩.pptx', 'VR 游戏设计答辩用演示文稿（含一份备份副本）。'),
    ('PPT制作指南.md', '答辩 PPT 制作规范与建议。'),
    ('slides_12_15_描述与文案.md', 'PPT 每页的详细描述与文案内容。'),
    ('missing_content_演讲稿.md', '答辩演讲稿，补充 PPT 中未展示的内容。'),
    ('arch_battle_system.svg', '战斗系统架构示意图（SVG 格式）。'),
    ('svg/', 'SVG 图表资源目录。'),
    ('extract_ppt.py', 'PPT 内容提取脚本（Python），用于从 .pptx 中提取文字。'),
    ('中期ppt.txt', '中期答辩 PPT 的文字内容存档。'),
    ('最后要求.txt', '项目最终交付要求记录。'),
]

table = doc.add_table(rows=1, cols=2, style='Light Grid Accent 1')
table.columns[0].width = Inches(2.5)
table.columns[1].width = Inches(4.2)
hdr = table.rows[0].cells
hdr[0].text = '文件名'
hdr[1].text = '说明'
for filename, desc in docs_files:
    row = table.add_row().cells
    row[0].text = filename
    row[1].text = desc

# ════════════════════════════════════════════════════════════════════════
# 7. 项目配置与解决方案
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('七、项目配置文件与解决方案', level=1)

cfg_items = [
    ('Test.sln', 'Visual Studio 解决方案文件，组织所有 C# 项目（Assembly-CSharp、PICO SDK 等）。'),
    ('Assembly-CSharp.csproj', '主 C# 脚本程序集项目文件，Assets/Script/ 中的脚本默认编译到此程序集。'),
    ('Assembly-CSharp-Editor.csproj', '编辑器脚本程序集项目文件（Editor 目录下的脚本）。'),
    ('PICO.Platform.csproj / PICO.TobSupport.csproj 等', 'PICO SDK 各模块的 C# 项目文件。'),
    ('Unity.XR.PICO*.csproj', 'PICO XR 插件相关项目文件。'),
    ('ps_driver_sdk.log', 'PICO 驱动 SDK 的日志文件。'),
    ('Assets/user.keystore', 'Android 签名密钥库，用于构建 Android APK。'),
    ('Assets/llms.txt', '指向 Claude/claude.ai 的 LLM 指南文件，帮助 AI 理解项目结构。'),
]

for name, desc in cfg_items:
    p = doc.add_paragraph()
    run = p.add_run(name + '  ')
    run.bold = True
    p.add_run(desc)

# ── footer ──────────────────────────────────────────────────────────────
doc.add_paragraph('')
doc.add_paragraph('— 文档结束 —').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── save ────────────────────────────────────────────────────────────────
out_path = 'e:/GameProject/7_VR_test/Test/Docs/readme.docx'
doc.save(out_path)
print(f'Done: {out_path}')
