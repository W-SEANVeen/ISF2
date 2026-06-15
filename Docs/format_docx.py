"""
Format the docx report:
- 题目: 小二号 (18pt), 加粗
- 一级标题: 小三号 (15pt), 加粗
- 二级标题: 四号 (14pt), 加粗
- 三级标题: 小四号 (12pt), 加粗
- 正文: 小四号 (12pt), 宋体, 1.5倍行距
- 英文/数字/罗马数字: Times New Roman
- 中文: 宋体 (SimSun)
"""

import re
import copy
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING

SRC = r"E:\GameProject\7_VR_test\Test\Docs\《孤城》——《综合挑战实践课II》2025-2026第2学期期末考试大作业报告.docx"
DST = r"E:\GameProject\7_VR_test\Test\Docs\《孤城》——《综合挑战实践课II》2025-2026第2学期期末考试大作业报告.docx"

# Font size constants (pt)
SIZE_TITLE  = 18   # 小二号
SIZE_H1     = 15   # 小三号
SIZE_H2     = 14   # 四号
SIZE_BODY   = 12   # 小四号

# ── helpers ──────────────────────────────────────────────────────

def set_run_font(run, size_pt=None, bold=None, latin_font='Times New Roman', ea_font='SimSun'):
    """Set Latin + East-Asian fonts on a run."""
    rPr = run._element.get_or_add_rPr()

    # build <w:rFonts>
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), latin_font)
    rFonts.set(qn('w:hAnsi'), latin_font)
    rFonts.set(qn('w:eastAsia'), ea_font)
    rFonts.set(qn('w:cs'), latin_font)

    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def set_paragraph_spacing(para, line_spacing=1.5):
    """Set 1.5x line spacing on a paragraph."""
    pf = para.paragraph_format
    pf.line_spacing = line_spacing


def process_paragraph(para, h_level='body'):
    """Apply formatting based on heading level."""
    if not para.text.strip():
        return

    for run in para.runs:
        if h_level == 'title':
            set_run_font(run, SIZE_TITLE, True)
        elif h_level == 'h1':
            set_run_font(run, SIZE_H1, True)
        elif h_level == 'h2':
            set_run_font(run, SIZE_H2, True)
        elif h_level == 'h3':
            set_run_font(run, SIZE_BODY, True)
        else:  # body
            set_run_font(run, SIZE_BODY, None)

    # Line spacing: 1.5 for body; keep headings' own spacing
    if h_level == 'body':
        set_paragraph_spacing(para, 1.5)

    # For body text, also set paragraph font (inheritance fallback)
    if h_level == 'body':
        pf = para.paragraph_format
        # Ensure runs without explicit size inherit
        pass


def classify_paragraph(para):
    """Determine formatting level for a paragraph."""
    text = para.text.strip()
    if not text:
        return None

    # ── Main title ──
    # The main project title at ~18pt
    style_name = para.style.name if para.style else ''
    if text.startswith('《孤城》'):
        return 'title'

    # ── Level-1 headings ──
    # "一、序论", "二、运行说明", "三、核心功能实现", "四、测试", "五、团队贡献度"
    if re.match(r'^[一二三四五六七八九十]、', text):
        return 'h1'
    if style_name == 'Heading 1':
        return 'h1'
    # Also h1 if it matches the pattern with 分 in it
    if re.match(r'^[一二三四五六七八九十]、.*分\)?$', text):
        return 'h1'

    # ── Level-3 headings (X.X.X) ──
    # "3.1.1 总体架构", "3.7.1 箭雨对象池"
    if re.match(r'^\d+\.\d+\.\d+\s', text):
        # Only treat as h3 if it's currently bold (heading-like)
        some_bold = any(r.font.bold for r in para.runs if r.font.bold is not None)
        if some_bold:
            return 'h3'

    # ── Level-2 headings (X.X) ──
    # "1.1 硬件环境", "2.1 运行前提", "3.7 性能优化", "4.1 PICO SDK..."
    if re.match(r'^\d+\.\d+\s', text):
        some_bold = any(r.font.bold for r in para.runs if r.font.bold is not None)
        if some_bold:
            return 'h2'

    # ── Special: 摘要 / 关键词 ──
    # Keep as body text

    return 'body'


def process_table_cells(table):
    """Apply body-text formatting to all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if not para.text.strip():
                    continue
                for run in para.runs:
                    set_run_font(run, SIZE_BODY, None)
                set_paragraph_spacing(para, 1.5)


# ── main ─────────────────────────────────────────────────────────

def main():
    doc = Document(SRC)

    print(f"Processing {len(doc.paragraphs)} paragraphs...")

    # ---- Process paragraphs ----
    for i, para in enumerate(doc.paragraphs):
        h_level = classify_paragraph(para)
        if h_level is None:
            continue

        # Special handling: cover page (before first heading)
        # We detect cover as paragraphs before "一、序论" that aren't the title
        # For these, keep existing sizes but still apply Times New Roman + SimSun
        text = para.text.strip()

        is_cover = False
        if i < 30 and text and h_level not in ('title', 'h1'):
            # Check if we're before the first h1
            first_h1_idx = None
            for j, pp in enumerate(doc.paragraphs):
                if classify_paragraph(pp) == 'h1':
                    first_h1_idx = j
                    break
            if first_h1_idx is not None and i < first_h1_idx:
                is_cover = True

        if is_cover:
            # Cover: don't change size/spacing, just set fonts
            for run in para.runs:
                set_run_font(run, size_pt=None, bold=None)
            continue

        if h_level == 'body':
            # Body text: 12pt, SimSun, 1.5x spacing, Latin=TNR
            for run in para.runs:
                set_run_font(run, SIZE_BODY, None)
            set_paragraph_spacing(para, 1.5)
        else:
            process_paragraph(para, h_level)

    # ---- Process tables ----
    print(f"Processing {len(doc.tables)} tables...")
    for table in doc.tables:
        process_table_cells(table)

    # ---- Save ----
    doc.save(DST)
    print(f"Done! Saved to {DST}")


if __name__ == '__main__':
    main()
