#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Check document formatting against requirements."""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from docx.shared import Pt, Emu
from docx.enum.text import WD_LINE_SPACING

doc_path = "《孤城》——《综合挑战实践课II》2025-2026第2学期期末考试大作业报告.docx"
doc = docx.Document(doc_path)

def pt_size(size_emu):
    if size_emu is None:
        return None
    return size_emu / 12700

def check_run_font(run, expected_cn_font=None, expected_en_font=None, expected_size=None, expected_bold=None):
    issues = []
    font = run.font

    if expected_size is not None:
        actual_size = pt_size(font.size) if font.size else None
        if actual_size and abs(actual_size - expected_size) > 0.5:
            issues.append(f"字号: 期望{expected_size}pt, 实际{actual_size:.1f}pt")
        elif actual_size is None:
            issues.append(f"字号: 未显式设置 (期望{expected_size}pt)")

    if expected_bold is not None:
        if font.bold != expected_bold:
            issues.append(f"加粗: 期望{expected_bold}, 实际{font.bold}")

    # Check rFonts element for East-Asian and ASCII fonts
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    try:
        rPr = run._element.find(f'.//{ns}rFonts')
        if rPr is not None:
            ea = rPr.get(f'{ns}eastAsia')
            ascii_f = rPr.get(f'{ns}ascii')
            h_ansi = rPr.get(f'{ns}hAnsi')

            if expected_cn_font and ea and ea != expected_cn_font:
                issues.append(f"东亚字体(eastAsia): 期望'{expected_cn_font}', 实际'{ea}'")
            elif expected_cn_font and not ea and font.name and font.name != expected_cn_font and font.name.lower() != 'times new roman':
                issues.append(f"字体名: 期望'{expected_cn_font}', 实际'{font.name}'")

            if expected_en_font:
                if ascii_f and ascii_f.lower() != expected_en_font.lower():
                    issues.append(f"ASCII字体: 期望'{expected_en_font}', 实际'{ascii_f}'")
                if h_ansi and h_ansi.lower() != expected_en_font.lower():
                    issues.append(f"ANSI字体(hAnsi): 期望'{expected_en_font}', 实际'{h_ansi}'")
    except Exception:
        pass

    return issues

def check_paragraph_spacing(para, expected_line_spacing=1.5):
    issues = []
    pf = para.paragraph_format
    spacing = pf.line_spacing
    spacing_rule = pf.line_spacing_rule

    if spacing_rule == WD_LINE_SPACING.MULTIPLE:
        if spacing and abs(spacing - expected_line_spacing) > 0.05:
            issues.append(f"行距: 期望{expected_line_spacing}倍, 实际{spacing}倍")
        elif spacing is None:
            issues.append(f"行距: MULTIPLE模式但值为None")
    elif spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE:
        pass
    elif spacing_rule == WD_LINE_SPACING.SINGLE:
        issues.append(f"行距: 单倍行距 (期望{expected_line_spacing}倍)")
    elif spacing_rule == WD_LINE_SPACING.DOUBLE:
        issues.append(f"行距: 双倍行距 (期望{expected_line_spacing}倍)")
    elif spacing_rule is None:
        issues.append(f"行距: 未显式设置 (期望{expected_line_spacing}倍)")
    else:
        issues.append(f"行距: 规则={spacing_rule}, 值={spacing}")
    return issues

def classify_paragraph(para, index):
    text = para.text.strip()
    if not text:
        return "空行"

    style_name = para.style.name if para.style else "无样式"

    max_size = 0
    for run in para.runs:
        if run.font.size:
            s = pt_size(run.font.size)
            if s and s > max_size:
                max_size = s

    if max_size and abs(max_size - 18) < 1:
        return "题目"
    elif max_size and abs(max_size - 16) < 1:
        return "三号(可能为一级标题)"
    elif max_size and abs(max_size - 15) < 1:
        return "一级标题"
    elif max_size and abs(max_size - 14) < 1:
        return "二级标题"
    elif max_size and abs(max_size - 12) < 1:
        return "正文"
    elif max_size and abs(max_size - 10.5) < 1:
        return "五号(可能为正文)"

    if style_name and ('heading' in style_name.lower() or '标题' in style_name or 'head' in style_name.lower()):
        return f"标题样式({style_name})"

    return f"未分类(最大字号{max_size:.0f}pt)" if max_size else "未分类(无字号)"

print("=" * 80)
print("文档格式检查报告")
print(f"文件: {doc_path}")
print("=" * 80)

total_issues = []
para_count = 0
category_stats = {}

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue

    para_count += 1
    classification = classify_paragraph(para, i)
    category_stats[classification] = category_stats.get(classification, 0) + 1

    para_issues = []

    if classification == "题目":
        for run in para.runs:
            para_issues.extend(check_run_font(run, "宋体", "Times New Roman", 18, True))

    elif classification in ("一级标题", "三号(可能为一级标题)"):
        expected_size = 15 if classification == "一级标题" else 16
        for run in para.runs:
            para_issues.extend(check_run_font(run, "宋体", "Times New Roman", expected_size, True))

    elif classification == "二级标题":
        for run in para.runs:
            para_issues.extend(check_run_font(run, "宋体", "Times New Roman", 14, True))

    elif classification in ("正文", "五号(可能为正文)"):
        expected_size = 12 if classification == "正文" else 10.5
        for run in para.runs:
            para_issues.extend(check_run_font(run, "宋体", "Times New Roman", expected_size, False))
        para_issues.extend(check_paragraph_spacing(para, 1.5))

    if para_issues:
        total_issues.append((i, classification, text[:60], para_issues))

print(f"\n共扫描 {para_count} 个非空段落")
print(f"\n段落类型分布:")
for cls, cnt in sorted(category_stats.items(), key=lambda x: -x[1]):
    print(f"  {cls}: {cnt} 段")

print(f"\n发现 {len(total_issues)} 个段落有格式问题:\n")
for idx, cls, txt, issues in total_issues:
    print(f"--- 段落 {idx} [{cls}] ---")
    print(f"  内容: {txt}")
    for iss in issues:
        print(f"  [问题] {iss}")
    print()

if not total_issues:
    print("[OK] 所有段落格式正确！")
else:
    print(f"[需要修正] 共 {len(total_issues)} 个段落需要修正")

# Detailed per-paragraph analysis
print("\n\n" + "=" * 80)
print("详细段落分析（全部）")
print("=" * 80)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue

    cls = classify_paragraph(para, i)
    print(f"\n--- 段落 {i} [{cls}] ---")
    print(f"  内容: {text[:80]}{'...' if len(text) > 80 else ''}")
    print(f"  样式: {para.style.name}")

    for j, run in enumerate(para.runs):
        if not run.text.strip():
            continue
        font = run.font
        size_str = f"{pt_size(font.size):.0f}pt" if font.size else "继承"
        bold_str = "加粗" if font.bold else "正常"

        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        ea = ""
        ascii_f = ""
        h_ansi = ""
        try:
            rPr = run._element.find(f'.//{ns}rFonts')
            if rPr is not None:
                ea = rPr.get(f'{ns}eastAsia') or ""
                ascii_f = rPr.get(f'{ns}ascii') or ""
                h_ansi = rPr.get(f'{ns}hAnsi') or ""
        except:
            pass

        run_text = run.text[:50].replace('\n', '\\n')
        print(f"    Run[{j}]: \"{run_text}\"")
        print(f"      字号={size_str} | {bold_str}")
        print(f"      font.name={font.name} | eastAsia={ea} | ascii={ascii_f} | hAnsi={h_ansi}")
